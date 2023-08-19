import streamlit as st
from io import BytesIO
from translator import text_translator
from docx import Document
from doc_reader import docx_parser

# Configure Streamlit page settings
st.set_page_config(page_title="Script Translator", page_icon=":page_with_curl:")

# Title of the web app
st.title("Script Translator")

# Taking inputs from the user
uploaded_file = st.file_uploader("Choose a docx file", type="docx")
target_language = st.selectbox("Select target language", ["English", "French", "Pirated English", "German"])
target_country = st.selectbox("Select target country", ["India", "United States", "Spain", "France", "Germany", "China"])

# Submit button
submit = st.button("Submit")

# Check if a file is uploaded and the submit button is clicked
if uploaded_file is not None and submit:
    # Parse the uploaded DOCX file
    text = docx_parser(uploaded_file)
    
    # Display the input text
    st.header("Your Input Text")
    st.write(text)

    # Translate the input text
    translated_text = text_translator(text, target_language, target_country)
    
    # Display the translated text
    st.header("Your Translated Text")
    st.write(translated_text)

    # Outputting in DOCX format
    output_doc = Document()
    output_doc.add_heading("Translated Text", level=1)
    output_doc.add_paragraph(translated_text)
    output_doc_stream = BytesIO()
    output_doc.save(output_doc_stream)
    output_doc_stream.seek(0)

    # Creating a download button for the translated text
    st.download_button(
        label="Download Translated Text (DOCX)",
        data=output_doc_stream,
        file_name="translated_text.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
